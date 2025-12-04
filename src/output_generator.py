import os
import re
import markdown
from datetime import datetime
from bs4 import BeautifulSoup

# Optional import for PDF generation
try:
    from xhtml2pdf import pisa
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    pisa = None

class OutputGenerator:
    def __init__(self, output_dir='outputs'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def _get_paths(self, prefix, ext):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{prefix}_{timestamp}.{ext}"
        return os.path.join(self.output_dir, filename)

    def generate_all_formats(self, content, prefix="report"):
        """Generate all file formats and return their paths/bytes"""
        return {
            'md': self.save_markdown(content, prefix),
            'txt': self.save_txt(content, prefix),
            'pdf': self.save_pdf(content, prefix),
            'docx': self.save_docx(content, prefix)
        }

    def save_markdown(self, content, prefix="report"):
        path = self._get_paths(prefix, "md")
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return path

    def save_txt(self, content, prefix="report"):
        path = self._get_paths(prefix, "txt")
        
        # Pre-clean markdown artifacts that BeautifulSoup might miss
        # Remove bold/italic markers
        clean_content = content.replace('**', '').replace('__', '')
        
        html = markdown.markdown(clean_content, extensions=['tables', 'fenced_code'])
        soup = BeautifulSoup(html, "html.parser")
        
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr']):
            element.append('\n')
            
        text_content = soup.get_text()
        text_content = re.sub(r'\n{3,}', '\n\n', text_content)
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(text_content.strip())
        return path
    
    def save_docx(self, content, prefix="report"):
        import docx
        from docx.shared import Pt, RGBColor
        
        path = self._get_paths(prefix, "docx")
        doc = docx.Document()
        
        lines = content.split('\n')
        in_code_block = False
        
        # Simple table tracking
        in_table = False
        table = None
        
        for line in lines:
            stripped = line.strip()
            
            # Code blocks
            if stripped.startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                p = doc.add_paragraph(line)
                p.style = 'Quote'
                continue
                
            # Table Detection
            if '|' in line and not in_code_block:
                # Check if it's a separator line
                if re.match(r'^\s*\|?[-:| ]+\|?\s*$', line):
                    continue
                
                # Split cells
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if not cells:
                    continue
                    
                if not in_table:
                    # Create new table with detected columns
                    table = doc.add_table(rows=0, cols=len(cells))
                    table.style = 'Table Grid'
                    in_table = True
                
                # Add row (ensure column count matches, or append to last row)
                try:
                    row_cells = table.add_row().cells
                    for i, cell_text in enumerate(cells):
                        if i < len(row_cells):
                            # Handle bold in table cells
                            clean_text = cell_text.replace('**', '').replace('__', '')
                            row_cells[i].text = clean_text
                except Exception:
                    # Fallback if columns don't match
                    pass
                continue
            else:
                in_table = False

            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # List items
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                self._add_run_with_formatting(p, text)
                
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(style='List Number')
                self._add_run_with_formatting(p, text)
                
            # Regular paragraph
            elif stripped:
                p = doc.add_paragraph()
                self._add_run_with_formatting(p, line)
                
        doc.save(path)
        return path

    def _add_run_with_formatting(self, paragraph, text):
        """Parse text for bold (**text**) and add runs to paragraph"""
        # Split by **...**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def save_pdf(self, content, prefix="report"):
        """
        Save content as PDF. Requires xhtml2pdf package.
        If not available, creates a placeholder file with instructions.
        """
        path = self._get_paths(prefix, "pdf")
        
        if not PDF_AVAILABLE:
            # Create a placeholder file with instructions
            placeholder_text = f"""
PDF generation requires xhtml2pdf package.

To install:
1. Install system dependencies: brew install cairo pkg-config
2. Install Python package: pip install xhtml2pdf

Original content (as text):
{content}
"""
            with open(path, "w", encoding='utf-8') as f:
                f.write(placeholder_text)
            return path
        
        # Clean emojis specifically for PDF generation to avoid black boxes
        # This regex removes characters in the emoji range
        # ranges: 1F600-1F64F (Emoticons), 1F300-1F5FF (Symbols & Pictographs), etc.
        # A simple broad regex for high unicode characters often catches emojis
        # Or we can just remove specific colored circles if we know them.
        # Safer approach: Remove common emoji ranges to ensure clean PDF.
        
        pdf_safe_content = re.sub(r'[^\x00-\x7F]+', '', content) if prefix == "status_report" else content
        
        html_content = markdown.markdown(pdf_safe_content, extensions=['tables', 'fenced_code'])
        
        styled_html = f"""
        <html>
        <head>
            <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
            <style>
                @page {{ size: A4; margin: 2cm; }}
                body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #333; }}
                h1 {{ color: #2c3e50; font-size: 18pt; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                h2 {{ color: #2c3e50; font-size: 14pt; margin-top: 20px; border-bottom: 1px solid #eee; }}
                
                /* Tables */
                table {{ width: 100%; border-collapse: collapse; margin-bottom: 15px; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; color: #333; }}
                
                code {{ background-color: #f8f9fa; padding: 2px 4px; font-family: Courier; }}
            </style>
        </head>
        <body>{html_content}</body>
        </html>
        """
        
        with open(path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(styled_html, dest=pdf_file, encoding='utf-8')
            
        if pisa_status.err:
            raise Exception("PDF generation failed")
        return path
