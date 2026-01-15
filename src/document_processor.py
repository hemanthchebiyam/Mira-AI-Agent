import os
import io
import re
import chardet
import pandas as pd
from typing import List, Optional, Tuple, Dict, Any

# PDF Parsers - try multiple options
PDF_PARSER = None
try:
    import pdfplumber
    PDF_PARSER = 'pdfplumber'
except ImportError:
    try:
        import PyPDF2
        PDF_PARSER = 'pypdf2'
    except ImportError:
        pass

# DOCX parser
try:
    import docx
except ImportError:
    docx = None

# LangChain imports for RAG
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
    def process_files(self, uploaded_files):
        """
        Process a list of uploaded files (Streamlit UploadedFile objects)
        and return a combined text string and file details.
        
        Args:
            uploaded_files: List of Streamlit UploadedFile objects
            
        Returns:
            dict with 'combined_text' and 'file_details' keys
        """
        combined_text = ""
        file_details = []

        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_ext = os.path.splitext(file_name)[1].lower()
            content = ""
            metadata = {}
            
            try:
                if file_ext == '.pdf':
                    content, metadata = self._read_pdf(uploaded_file)
                elif file_ext == '.docx':
                    content, metadata = self._read_docx(uploaded_file)
                elif file_ext == '.txt':
                    content, metadata = self._read_txt(uploaded_file)
                elif file_ext in ['.xlsx', '.xls']:
                    content, metadata = self._read_excel(uploaded_file)
                else:
                    content = f"[Skipped unsupported file format: {file_name}]"
                    metadata = {"error": "Unsupported format"}

                if content and not content.startswith("[Skipped") and not content.startswith("Error"):
                    # Add metadata info if available
                    metadata_str = ""
                    if metadata:
                        if 'pages' in metadata:
                            metadata_str = f" ({metadata['pages']} pages)"
                        elif 'sheets' in metadata:
                            metadata_str = f" ({len(metadata['sheets'])} sheets)"
                        elif 'encoding' in metadata:
                            metadata_str = f" (encoding: {metadata['encoding']})"
                    
                    combined_text += f"\n\n--- Content from {file_name}{metadata_str} ---\n{content}"
                    file_details.append({
                        "name": file_name, 
                        "status": "Processed",
                        "metadata": metadata
                    })
                else:
                    file_details.append({
                        "name": file_name, 
                        "status": "Empty or Unreadable",
                        "metadata": metadata
                    })

            except Exception as e:
                error_msg = str(e)
                file_details.append({
                    "name": file_name, 
                    "status": f"Error: {error_msg}",
                    "metadata": {"error": error_msg}
                })
                combined_text += f"\n\n--- Error reading {file_name} ---\n{error_msg}"

        return {
            "combined_text": combined_text.strip(),
            "file_details": file_details
        }

    def _read_pdf(self, file) -> Tuple[str, Dict[str, Any]]:
        """
        Read PDF file with improved text extraction.
        Tries pdfplumber first, then PyPDF2 as fallback.
        
        Returns:
            tuple: (text_content, metadata_dict)
        """
        text = ""
        metadata = {}
        
        try:
            # Reset file pointer
            file.seek(0)
            
            if PDF_PARSER == 'pdfplumber':
                import pdfplumber
                with pdfplumber.open(file) as pdf:
                    metadata['pages'] = len(pdf.pages)
                    metadata['parser'] = 'pdfplumber'
                    
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n[Page {i+1}]\n{page_text}\n"
                        
                        # Try to extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                text += "\n[Table]\n"
                                for row in table:
                                    if row:
                                        text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                                text += "\n"
            
            elif PDF_PARSER == 'pypdf2':
                import PyPDF2
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                metadata['parser'] = 'pypdf2'
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n[Page {i+1}]\n{page_text}\n"
            
            else:
                return "Error: No PDF parser available. Install pdfplumber or PyPDF2.", {"error": "No PDF parser"}
            
            # Clean up excessive whitespace
            text = re.sub(r'\n{3,}', '\n\n', text)
            
        except Exception as e:
            return f"Error reading PDF: {str(e)}", {"error": str(e)}
        
        return text.strip(), metadata

    def _read_docx(self, file) -> Tuple[str, Dict[str, Any]]:
        """
        Read DOCX file with support for paragraphs, tables, lists, and formatting.
        
        Returns:
            tuple: (text_content, metadata_dict)
        """
        text = ""
        metadata = {}
        
        if docx is None:
            return "Error: python-docx not installed", {"error": "python-docx not installed"}
        
        try:
            file.seek(0)
            doc = docx.Document(file)
            
            # Count elements
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            metadata['paragraphs'] = para_count
            metadata['tables'] = table_count
            
            # Process paragraphs
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        text += f"\n{'#' * int(level)} {para_text}\n" if level.isdigit() else f"\n## {para_text}\n"
                    else:
                        text += para_text + "\n"
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                text += f"\n[Table {table_idx + 1}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            # Clean up
            text = re.sub(r'\n{3,}', '\n\n', text)
            
        except Exception as e:
            return f"Error reading DOCX: {str(e)}", {"error": str(e)}
        
        return text.strip(), metadata

    def _read_txt(self, file) -> Tuple[str, Dict[str, Any]]:
        """
        Read TXT file with automatic encoding detection.
        
        Returns:
            tuple: (text_content, metadata_dict)
        """
        metadata = {}
        
        try:
            file.seek(0)
            raw_data = file.read()
            
            # Detect encoding
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            metadata['encoding'] = encoding
            metadata['confidence'] = f"{confidence:.2%}"
            
            # Try to decode with detected encoding
            try:
                text = raw_data.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                text = raw_data.decode('utf-8', errors='replace')
                metadata['encoding'] = 'utf-8 (fallback)'
            
            # Clean up line endings
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            
        except Exception as e:
            return f"Error reading TXT: {str(e)}", {"error": str(e)}
        
        return text.strip(), metadata

    def _read_excel(self, file) -> Tuple[str, Dict[str, Any]]:
        """
        Read Excel file with improved structure preservation.
        Handles multiple sheets, preserves table structure, and extracts metadata.
        
        Returns:
            tuple: (text_content, metadata_dict)
        """
        text = ""
        metadata = {}
        
        try:
            file.seek(0)
            
            # Determine engine based on file extension
            file_ext = getattr(file, 'name', '').lower()
            engine = 'openpyxl' if file_ext.endswith('.xlsx') else None
            
            # Read all sheets
            excel_file = pd.ExcelFile(file, engine=engine)
            sheet_names = excel_file.sheet_names
            metadata['sheets'] = sheet_names
            metadata['sheet_count'] = len(sheet_names)
            
            for sheet_name in sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                text += f"\n{'='*60}\n"
                text += f"Sheet: {sheet_name}\n"
                text += f"{'='*60}\n\n"
                
                if df.empty:
                    text += "[Empty sheet]\n\n"
                    continue
                
                # Add column names
                text += "Columns: " + ", ".join(str(col) for col in df.columns) + "\n\n"
                
                # Convert DataFrame to string with better formatting
                if len(df) <= 100:
                    text += df.to_string(index=False, max_rows=None) + "\n\n"
                else:
                    text += "[First 50 rows]\n"
                    text += df.head(50).to_string(index=False) + "\n\n"
                    text += f"[... {len(df) - 100} rows omitted ...]\n\n"
                    text += "[Last 50 rows]\n"
                    text += df.tail(50).to_string(index=False) + "\n\n"
                
                # Add summary statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    text += "\n[Summary Statistics]\n"
                    summary = df[numeric_cols].describe()
                    text += summary.to_string() + "\n\n"
            
            excel_file.close()
            
        except Exception as e:
            return f"Error reading Excel: {str(e)}", {"error": str(e)}
        
        return text.strip(), metadata

    def process_files_to_documents(self, uploaded_files) -> List[Document]:
        """
        Process uploaded files and return LangChain Document objects.
        Each document chunk includes metadata about the source file.
        """
        all_documents = []
        
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_ext = os.path.splitext(file_name)[1].lower()
            content = ""
            
            try:
                # Reset file pointer before reading
                uploaded_file.seek(0)
                
                # All read methods return (content, metadata) tuple
                if file_ext == '.pdf':
                    content, _ = self._read_pdf(uploaded_file)
                elif file_ext == '.docx':
                    content, _ = self._read_docx(uploaded_file)
                elif file_ext == '.txt':
                    content, _ = self._read_txt(uploaded_file)
                elif file_ext in ['.xlsx', '.xls']:
                    content, _ = self._read_excel(uploaded_file)
                else:
                    continue  # Skip unsupported formats
                
                if content and content.strip() and not content.startswith("Error"):
                    # Split content into chunks
                    chunks = self.text_splitter.split_text(content)
                    
                    # Create Document objects with metadata
                    for i, chunk in enumerate(chunks):
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                "source": file_name,
                                "chunk_index": i,
                                "total_chunks": len(chunks)
                            }
                        )
                        all_documents.append(doc)
                        
            except Exception as e:
                # Log error but continue processing other files
                print(f"Error processing {file_name}: {str(e)}")
                continue
        
        return all_documents

    def create_vector_store_simple(self, uploaded_files, api_key: str):
        """
        Create a simple in-memory vector store using FAISS (more stable than ChromaDB).
        
        Args:
            uploaded_files: List of Streamlit UploadedFile objects
            api_key: OpenAI API key for embeddings
            
        Returns:
            FAISS vector store instance or None if no documents
        """
        # Process files into LangChain Documents
        documents = self.process_files_to_documents(uploaded_files)
        
        if not documents:
            return None
        
        # Create embeddings
        embeddings = OpenAIEmbeddings(api_key=api_key)
        
        # Create FAISS vector store (simpler, no external dependencies)
        vector_store = FAISS.from_documents(
            documents=documents,
            embedding=embeddings
        )
        
        return vector_store
    
    def get_combined_text_from_documents(self, documents: List[Document]) -> str:
        """
        Combine all document chunks back into a single text string.
        Useful for the Project Plan generation that needs full context.
        """
        if not documents:
            return ""
        
        # Group by source file to maintain structure
        sources = {}
        for doc in documents:
            source = doc.metadata.get("source", "Unknown")
            if source not in sources:
                sources[source] = []
            sources[source].append(doc.page_content)
        
        # Combine with source headers
        combined = ""
        for source, chunks in sources.items():
            combined += f"\n\n--- Content from {source} ---\n"
            combined += "\n".join(chunks)
        
        return combined.strip()
